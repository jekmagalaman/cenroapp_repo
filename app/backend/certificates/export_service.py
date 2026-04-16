from io import BytesIO
from pathlib import Path

from .models import Certificate


_APPROVED_BY = 'CARDELAR STEVIE ANGEL M. MADRIÑAN'
_APPROVED_TITLE_1 = 'CG Assistant Department Head II'
_APPROVED_TITLE_2 = '(Assistant City ENRO)'
_MARINE_TEMPLATE = (
  Path(__file__).resolve().parent / 'docx_templates' / 'MARINE CERTIFICATION.docx'
)


def _is_marine(cert: Certificate) -> bool:
  return cert.certificate_type == 'marine_certification'


def _checked(cert: Certificate, expected: str) -> str:
  return 'x' if (cert.license_type or '').strip().lower() == expected.lower() else ' '


def _fill_marine_template_docx(doc, cert: Certificate) -> None:
  for p in doc.paragraphs:
    text = p.text or ''
    if 'THIS IS TO CERTIFY that Mr. / Ms.' in text:
      p.text = (
        f'THIS IS TO CERTIFY that Mr. / Ms. {cert.applicant_name} of '
        f'{cert.applicant_address}, Puerto Princesa City has been inspected by Bantay '
        "Dagat and hereby endorsing his/her request for approval and issuance of Mayor's "
        'Permit to operate/engage in the business of:'
      )
    elif '(  ) New' in text and 'Renew' in text:
      p.text = (
        f'({ _checked(cert, "New") }) New'
        f'            ({ _checked(cert, "Renew") }) Renew'
      )
    elif 'NATURE OF BUSINESS' in text:
      p.text = f'NATURE OF BUSINESS    :    {cert.nature_of_business}'
    elif 'BUSINESS ADDRESS' in text:
      p.text = f'BUSINESS ADDRESS    :    {cert.business_address}'
    elif 'NAME OF BUSINESS' in text:
      p.text = f'NAME OF BUSINESS    :    {cert.business_name}'
    elif 'Issued this' in text:
      p.text = f'Issued this {cert.issued_date}.'
    elif 'Contact No' in text:
      p.text = f'Contact No    :    {cert.contact_number}'
    elif 'Inspected by:' in text and 'Approved by:' in text:
      p.text = f'Inspected by: {cert.inspector_name}                    Approved by:'


def _rows(cert: Certificate):
  return [
    ('Control Number', cert.control_number),
    ('Certificate Type', cert.get_certificate_type_display()),
    ('Applicant Name', cert.applicant_name),
    ('Applicant Address', cert.applicant_address),
    ('License Type', cert.license_type),
    ('Nature of Business', cert.nature_of_business),
    ('Business Name', cert.business_name),
    ('Business Address', cert.business_address),
    ('Contact Number', cert.contact_number),
    ('Issued Date', cert.issued_date),
    ('Inspector Name', cert.inspector_name),
    ('Sent By Account', cert.created_by.username if cert.created_by else ''),
    ('Created At', cert.created_at.strftime('%Y-%m-%d %H:%M:%S')),
  ]


def build_pdf_bytes(cert: Certificate) -> bytes:
  try:
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfgen import canvas
  except Exception as exc:  # pragma: no cover
    raise RuntimeError('PDF export dependency missing: install reportlab') from exc

  buf = BytesIO()
  c = canvas.Canvas(buf, pagesize=A4)
  _width, height = A4

  if _is_marine(cert):
    y = height - 44
    c.setFont('Helvetica', 10)
    c.drawString(52, y, cert.control_number)
    y -= 16

    c.setFont('Helvetica-Bold', 11)
    c.drawCentredString(298, y, 'Republic of the Philippines')
    y -= 14
    c.drawCentredString(298, y, 'City Government of Puerto Princesa')
    y -= 14
    c.drawCentredString(
      298, y, 'Office of the City Environment and Natural Resources Officer'
    )
    y -= 14
    c.drawCentredString(298, y, 'ENVIRONMENTAL LAW ENFORCEMENT DIVISION (ELED)')
    y -= 14
    c.drawCentredString(298, y, 'Bantay Dagat Section')
    y -= 14
    c.setFont('Helvetica', 9)
    c.drawCentredString(
      298,
      y,
      'Ground Floor, Old City Hall Building, Bgy. Sta. Monica, Puerto Princesa City',
    )
    y -= 12
    c.drawCentredString(298, y, 'bantaydagat.ppc@gmail.com')

    y -= 24
    c.setFont('Helvetica-Bold', 14)
    c.drawCentredString(298, y, 'CERTIFICATE OF INSPECTION')
    y -= 14
    c.setFont('Helvetica', 11)
    c.drawCentredString(298, y, '(Marine Products)')

    y -= 26
    c.setFont('Helvetica-Bold', 11)
    c.drawString(52, y, 'TO WHOM IT MAY CONCERN:')
    y -= 18
    c.setFont('Helvetica', 10)
    c.drawString(
      52, y, f'THIS IS TO CERTIFY that Mr. / Ms. {cert.applicant_name}'
    )
    y -= 14
    c.drawString(
      52,
      y,
      f'of {cert.applicant_address}, Puerto Princesa City has been inspected by Bantay Dagat',
    )
    y -= 14
    c.drawString(
      52,
      y,
      "and hereby endorsing his/her request for approval and issuance of Mayor's Permit",
    )
    y -= 14
    c.drawString(52, y, 'to operate/engage in the business of:')
    y -= 20
    c.drawString(
      52,
      y,
      f'({ _checked(cert, "New") }) New     ({ _checked(cert, "Renew") }) Renew',
    )
    y -= 16
    c.drawString(52, y, 'KIND OF LICENSE APPLIED : BUSINESS PERMIT')
    y -= 14
    c.drawString(52, y, f'NATURE OF BUSINESS : {cert.nature_of_business}')
    y -= 14
    c.drawString(52, y, f'BUSINESS ADDRESS : {cert.business_address}')
    y -= 14
    c.drawString(52, y, f'NAME OF BUSINESS : {cert.business_name}')
    y -= 20
    c.drawString(
      52,
      y,
      'FURTHER CERTIFY that the above described business including the proposed location or area,',
    )
    y -= 14
    c.drawString(
      52,
      y,
      'prior to approval, does not pose any destruction/obstruction to the Ecological and Marine',
    )
    y -= 14
    c.drawString(52, y, 'Resources of the City.')
    y -= 20
    c.drawString(52, y, f'Issued this {cert.issued_date}.')
    y -= 20
    c.drawString(52, y, 'Conformed :')
    y -= 24
    c.drawString(52, y, '_____________________________')
    y -= 12
    c.drawString(52, y, '(Signature over Printed Name)')
    y -= 12
    c.drawString(52, y, 'Owner/Representative')
    y -= 14
    c.drawString(52, y, f'Contact No : {cert.contact_number}')

    c.drawString(330, y + 14, f'Inspected by: {cert.inspector_name}')
    c.drawString(330, y - 6, 'Approved by:')
    c.drawString(330, y - 20, 'BY AUTHORITY OF THE CITY ENRO:')
    c.drawString(330, y - 38, '___________________________')
    c.setFont('Helvetica-Bold', 10)
    c.drawString(330, y - 52, _APPROVED_BY)
    c.setFont('Helvetica', 10)
    c.drawString(330, y - 66, _APPROVED_TITLE_1)
    c.drawString(330, y - 80, _APPROVED_TITLE_2)
  else:
    y = height - 50
    c.setFont('Helvetica-Bold', 14)
    c.drawString(50, y, 'Certificate Export')
    y -= 28
    c.setFont('Helvetica', 10)

    for label, value in _rows(cert):
      text = f'{label}: {value or ""}'
      if len(text) > 130:
        text = text[:127] + '...'
      c.drawString(50, y, text)
      y -= 16
      if y < 50:
        c.showPage()
        c.setFont('Helvetica', 10)
        y = height - 50

  c.save()
  return buf.getvalue()


def build_excel_bytes(cert: Certificate) -> bytes:
  try:
    from openpyxl import Workbook
    from openpyxl.styles import Font
  except Exception as exc:  # pragma: no cover
    raise RuntimeError('Excel export dependency missing: install openpyxl') from exc

  wb = Workbook()
  ws = wb.active
  ws.title = 'Marine Certificate' if _is_marine(cert) else 'Certificate'

  if _is_marine(cert):
    ws['A1'] = cert.control_number
    ws['A3'] = 'Republic of the Philippines'
    ws['A4'] = 'City Government of Puerto Princesa'
    ws['A5'] = 'Office of the City Environment and Natural Resources Officer'
    ws['A6'] = 'ENVIRONMENTAL LAW ENFORCEMENT DIVISION (ELED)'
    ws['A7'] = 'Bantay Dagat Section'
    ws['A9'] = 'CERTIFICATE OF INSPECTION'
    ws['A10'] = '(Marine Products)'
    ws['A12'] = 'TO WHOM IT MAY CONCERN:'
    ws['A13'] = (
      f'THIS IS TO CERTIFY that Mr./Ms. {cert.applicant_name} of '
      f'{cert.applicant_address}, Puerto Princesa City has been inspected by Bantay Dagat.'
    )
    ws['A14'] = "Permit to operate/engage in the business of:"
    ws['A15'] = f'({ _checked(cert, "New") }) New    ({ _checked(cert, "Renew") }) Renew'
    ws['A16'] = 'KIND OF LICENSE APPLIED : BUSINESS PERMIT'
    ws['A17'] = f'NATURE OF BUSINESS : {cert.nature_of_business}'
    ws['A18'] = f'BUSINESS ADDRESS : {cert.business_address}'
    ws['A19'] = f'NAME OF BUSINESS : {cert.business_name}'
    ws['A21'] = (
      'FURTHER CERTIFY that the above described business does not pose '
      'destruction/obstruction to ecological and marine resources.'
    )
    ws['A22'] = f'Issued this {cert.issued_date}.'
    ws['A24'] = 'Conformed :'
    ws['A25'] = '_____________________________'
    ws['A26'] = '(Signature over Printed Name) Owner/Representative'
    ws['A27'] = f'Contact No : {cert.contact_number}'
    ws['A29'] = f'Inspected by: {cert.inspector_name}'
    ws['A30'] = 'Approved by:'
    ws['A31'] = 'BY AUTHORITY OF THE CITY ENRO:'
    ws['A32'] = '___________________________'
    ws['A33'] = _APPROVED_BY
    ws['A34'] = _APPROVED_TITLE_1
    ws['A35'] = _APPROVED_TITLE_2
    for cell in ('A3', 'A4', 'A5', 'A6', 'A7', 'A9', 'A12', 'A33'):
      ws[cell].font = Font(bold=True)
    ws.column_dimensions['A'].width = 120
  else:
    ws['A1'] = 'Certificate Export'
    ws['A1'].font = Font(bold=True, size=14)
    row = 3
    for label, value in _rows(cert):
      ws[f'A{row}'] = label
      ws[f'A{row}'].font = Font(bold=True)
      ws[f'B{row}'] = value or ''
      row += 1
    ws.column_dimensions['A'].width = 28
    ws.column_dimensions['B'].width = 60

  buf = BytesIO()
  wb.save(buf)
  return buf.getvalue()


def build_word_bytes(cert: Certificate) -> bytes:
  try:
    from docx import Document
  except Exception as exc:  # pragma: no cover
    raise RuntimeError('Word export dependency missing: install python-docx') from exc

  if _is_marine(cert) and _MARINE_TEMPLATE.exists():
    doc = Document(str(_MARINE_TEMPLATE))
    _fill_marine_template_docx(doc, cert)
  else:
    doc = Document()
    doc.add_heading('Certificate Export', level=1)
    for label, value in _rows(cert):
      p = doc.add_paragraph()
      p.add_run(f'{label}: ').bold = True
      p.add_run(value or '')

  buf = BytesIO()
  doc.save(buf)
  return buf.getvalue()
